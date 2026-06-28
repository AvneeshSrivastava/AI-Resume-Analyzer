"""
DOCX document extractor.

Provides DOCX text extraction using the python-docx library.
"""

from io import BytesIO

from docx import Document
from fastapi import UploadFile

from ai_resume_analyzer.exceptions.document import CorruptedDocumentException
from ai_resume_analyzer.services.extractors.base import BaseDocumentExtractor


class DocxExtractor(BaseDocumentExtractor):
    """
    Extracts text from DOCX documents.
    """

    async def extract_text(self, file: UploadFile) -> str:
        """
        Extract text from a DOCX document.

        Args:
            file: Uploaded DOCX file.

        Returns:
            Extracted text from all paragraphs.
        """

        docx_bytes = await file.read()

        document = Document(BytesIO(docx_bytes))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs)

    async def validate_document(self, file: UploadFile) -> None:
        """
        Validate the uploaded DOCX document.

        Args:
            file: Uploaded DOCX file.

        Raises:
            CorruptedDocumentException:
                If the document is corrupted or unreadable.
        """

        try:
            file.file.seek(0)

            Document(file.file)

        except Exception as exc:
            raise CorruptedDocumentException() from exc

        finally:
            # Reset stream for extraction
            file.file.seek(0)
